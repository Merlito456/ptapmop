import os
import re
import io
import base64
from copy import deepcopy
import streamlit as st

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from streamlit_js_eval import streamlit_js_eval


TEMPLATE_FILE = "Template.docx"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

MODE_SINGLE      = "Single Tapping Point"
MODE_SAME_RS     = "Main + RDNT – Same Rectifier"
MODE_SEPARATE_RS = "Main + RDNT – Separate Rectifiers"


# =============================================================
# DOCX HELPERS
# =============================================================

def normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def safe_filename(s: str) -> str:
    s = re.sub(r'[\\/*?:"<>|]+', "_", s)
    return s.strip().replace(" ", "_")


def iter_all_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def paragraph_full_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def set_paragraph_text(paragraph, text: str):
    for i, run in enumerate(paragraph.runs):
        run.text = text if i == 0 else ""
    if not paragraph.runs:
        paragraph.add_run(text)


def replace_text_in_paragraph_preserve_format(paragraph, replacements: dict):
    full = paragraph_full_text(paragraph)
    if not full.strip():
        return
    new_full = full
    changed = False
    for old, new in replacements.items():
        if old and old in new_full:
            new_full = new_full.replace(old, new)
            changed = True
    if not changed:
        return
    if paragraph.runs:
        first_run = paragraph.runs[0]
        bold      = first_run.bold
        italic    = first_run.italic
        underline = first_run.underline
        font_name = first_run.font.name
        font_size = first_run.font.size
        try:
            font_color = (
                first_run.font.color.rgb
                if first_run.font.color and first_run.font.color.type
                else None
            )
        except Exception:
            font_color = None
        for run in paragraph.runs:
            run.text = ""
        first_run.text      = new_full
        first_run.bold      = bold
        first_run.italic    = italic
        first_run.underline = underline
        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size
        if font_color:
            try:
                first_run.font.color.rgb = font_color
            except Exception:
                pass
    else:
        paragraph.add_run(new_full)


def replace_text_in_container(container, replacements: dict):
    for p in iter_all_paragraphs(container):
        replace_text_in_paragraph_preserve_format(p, replacements)


def replace_in_xml_text_nodes(xml_element, replacements: dict):
    ns = WORD_NS
    for t_node in xml_element.iter(f"{{{ns}}}t"):
        if t_node.text:
            new_text = t_node.text
            for old, new in replacements.items():
                if old and old in new_text:
                    new_text = new_text.replace(old, new)
            if new_text != t_node.text:
                t_node.text = new_text


def replace_everywhere(doc, replacements: dict):
    replace_text_in_container(doc, replacements)
    for section in doc.sections:
        replace_text_in_container(section.header, replacements)
        replace_text_in_container(section.footer, replacements)
        replace_in_xml_text_nodes(section.header._element, replacements)
        replace_in_xml_text_nodes(section.footer._element, replacements)
        try:
            replace_text_in_container(section.even_page_header, replacements)
            replace_in_xml_text_nodes(
                section.even_page_header._element, replacements)
        except Exception:
            pass
        try:
            replace_text_in_container(section.first_page_header, replacements)
            replace_in_xml_text_nodes(
                section.first_page_header._element, replacements)
        except Exception:
            pass


def find_paragraph_containing(container, needles, case_insensitive=True):
    for p in iter_all_paragraphs(container):
        txt = paragraph_full_text(p)
        chk = txt.lower() if case_insensitive else txt
        for needle in needles:
            nd = needle.lower() if case_insensitive else needle
            if nd in chk:
                return p
    return None


def find_in_doc(doc, needles):
    p = find_paragraph_containing(doc, needles)
    if p:
        return p, "body"
    for i, section in enumerate(doc.sections):
        p = find_paragraph_containing(section.header, needles)
        if p:
            return p, f"header_{i}"
        p = find_paragraph_containing(section.footer, needles)
        if p:
            return p, f"footer_{i}"
    return None, None


def insert_paragraph_after(ref_paragraph, text=""):
    """Plain insert — no formatting copy."""
    new_para = ref_paragraph._parent.add_paragraph()
    ref_paragraph._p.addnext(new_para._p)
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_after_copy_format(ref_paragraph, text=""):
    """
    Insert a new paragraph after ref_paragraph and copy its
    paragraph properties (indentation, spacing, alignment) and
    first run's character properties (font, size, bold, italic,
    underline, color) so the new line matches the reference exactly.
    """
    new_para = ref_paragraph._parent.add_paragraph()
    ref_paragraph._p.addnext(new_para._p)

    # Copy paragraph-level properties (pPr)
    ref_pPr = ref_paragraph._p.pPr
    if ref_pPr is not None:
        new_pPr = deepcopy(ref_pPr)
        # Replace or insert pPr in new paragraph
        existing_pPr = new_para._p.pPr
        if existing_pPr is not None:
            new_para._p.remove(existing_pPr)
        new_para._p.insert(0, new_pPr)

    if text:
        new_run = new_para.add_run(text)
        # Copy first run's character properties
        if ref_paragraph.runs:
            ref_run = ref_paragraph.runs[0]
            try:
                new_run.bold      = ref_run.bold
                new_run.italic    = ref_run.italic
                new_run.underline = ref_run.underline
                if ref_run.font.name:
                    new_run.font.name = ref_run.font.name
                if ref_run.font.size:
                    new_run.font.size = ref_run.font.size
                if ref_run.font.color and ref_run.font.color.type:
                    new_run.font.color.rgb = ref_run.font.color.rgb
            except Exception:
                pass

    return new_para


def replace_placeholder_with_image(doc, placeholders, image_bytes,
                                    width=Inches(5.0)):
    for p in iter_all_paragraphs(doc):
        full = paragraph_full_text(p)
        for ph in placeholders:
            if ph in full:
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].add_picture(
                        io.BytesIO(image_bytes), width=width)
                else:
                    p.add_run().add_picture(
                        io.BytesIO(image_bytes), width=width)
                return ph
    return None


def clear_placeholders(doc, placeholders):
    replace_everywhere(doc, {ph: "" for ph in placeholders})


# =============================================================
# AUDIT
# =============================================================

def audit_remaining_mf2(doc) -> list:
    suspects = ["MF-2", "MF2", "OLT MF", "Lightspan"]
    found = []
    for p in iter_all_paragraphs(doc):
        txt = paragraph_full_text(p)
        for s in suspects:
            if s in txt:
                found.append(f"[body/table] {txt.strip()[:120]}")
                break
    for section in doc.sections:
        ns = WORD_NS
        for t_node in section.header._element.iter(f"{{{ns}}}t"):
            txt = t_node.text or ""
            for s in suspects:
                if s in txt:
                    found.append(f"[header-xml] {txt.strip()[:120]}")
                    break
    return found


# =============================================================
# DEBUG
# =============================================================

def debug_list_all_text(doc) -> list:
    lines = []
    for p in iter_all_paragraphs(doc):
        txt = paragraph_full_text(p).strip()
        if txt:
            lines.append(f"[BODY] {txt}")
    for i, section in enumerate(doc.sections):
        for p in iter_all_paragraphs(section.header):
            txt = paragraph_full_text(p).strip()
            if txt:
                lines.append(f"[HEADER s{i}] {txt}")
        for p in iter_all_paragraphs(section.footer):
            txt = paragraph_full_text(p).strip()
            if txt:
                lines.append(f"[FOOTER s{i}] {txt}")
        ns = WORD_NS
        for t_node in section.header._element.iter(f"{{{ns}}}t"):
            txt = (t_node.text or "").strip()
            if txt:
                lines.append(f"[HEADER-XML s{i}] {txt}")
        for t_node in section.footer._element.iter(f"{{{ns}}}t"):
            txt = (t_node.text or "").strip()
            if txt:
                lines.append(f"[FOOTER-XML s{i}] {txt}")
    return lines


# =============================================================
# PLACEHOLDER VARIANTS
# =============================================================

RS1_LOAD_PH = [
    "{{RS1 Load schedule image}}",
    "{{RS1 Load schedule image}",
    "{{RS1 load schedule image}}",
    "{{RS1 load schedule image}",
]
RS2_LOAD_PH = [
    "{{RS2 Load schedule image}}",
    "{{RS2 Load schedule image}",
    "{{RS2 load schedule image}}",
    "{{RS2 load schedule image}",
]
RS1_EXIST_PH = [
    "{{RS1 EXISTING IMAGE}}",
    "{{RS1 EXISTING IMAGE}",
    "{{RS1 existing image}}",
    "{{RS1 existing image}",
]
RS2_EXIST_PH = [
    "{{RS2 EXISTING IMAGE}}",
    "{{RS2 EXISTING IMAGE}",
    "{{RS2 existing image}}",
    "{{RS2 existing image}",
]


# =============================================================
# BUSINESS LOGIC
# =============================================================

def build_olt_label(equipment: str, custom_olt_label: str) -> str:
    if normalize_spaces(equipment).lower() == "nokia lightspan mf-2":
        return "OLT MF-2"
    return normalize_spaces(custom_olt_label) or normalize_spaces(equipment)


def build_replacements(data: dict) -> dict:
    equipment = data["equipment"]
    olt_label = build_olt_label(equipment, data["olt_label_custom"])
    site_name = data["site_name"]
    plaid     = data["plaid"]

    parts  = equipment.strip().split()
    vendor = parts[0] if parts else "Nokia"
    model  = parts[-1] if len(parts) > 1 else equipment

    return {
        "Nokia Lightspan MF-2":     equipment,
        "Nokia Lightspan MF2":      equipment,
        "Lightspan MF-2":           equipment,
        "Lightspan MF2":            equipment,
        "Nokia / OLT MF-2":         f"{vendor} / {olt_label}",
        "Nokia/ OLT MF-2":          f"{vendor} / {olt_label}",
        "Nokia /OLT MF-2":          f"{vendor} / {olt_label}",
        "Nokia OLT MF-2":           f"{vendor} {olt_label}",
        "Nokia OLT MF2":            f"{vendor} {olt_label}",
        "OLT MF-2":                 olt_label,
        "OLT MF2":                  olt_label,
        "MF-2":                     model,
        "MF2":                      model,
        "CDO-604_MIN995":           f"{site_name}_{plaid}",
        "CDO-604_MIN699":           f"{site_name}_{plaid}",
        "CDO-604":                  site_name,
        "MIN699":                   plaid,
        "MIN995":                   plaid,
        "John Carlo Rabanes":       data["prepared_by"],
        "OLT Rollout Engineer":     data["position"],
        "OLT Engineer":             data["position"],
        "< May 19- June 19, 2026 10:00AM-6:00PM>": data["target_datetime"],
        "May 19- June 19, 2026 10:00AM-6:00PM":    data["target_datetime"],
        "{{SITE_NAME}}":            site_name,
        "{{PLAID}}":                plaid,
        "{{EQUIPMENT}}":            equipment,
        "{{PREPARED_BY}}":          data["prepared_by"],
        "{{POSITION}}":             data["position"],
        "{{TARGET_DATETIME}}":      data["target_datetime"],
    }


def fuse_no_line(load: str, olt_label: str, equipment: str,
                  label: str = None) -> str:
    """Build a FUSE No line string."""
    txt = (
        f"FUSE No: {load} {olt_label} "
        f"– ({normalize_spaces(equipment)} Power tapping point)"
    )
    if label:
        txt += f" – {label}"
    return txt


def build_tapping_summary(tapping_mode: str, rs_entries: list,
                           equipment: str) -> list:
    """Build planned activity fuse assignment lines."""
    lines = []

    if tapping_mode == MODE_SINGLE:
        rs = rs_entries[0] if rs_entries else None
        if rs and rs.get("load"):
            line = f"Fuse {rs['load']}"
            if rs.get("ampere"):
                line += f" ({rs['ampere']})"
            line += f" : {equipment} Power Tapped"
            if rs.get("name"):
                line += f"  [{rs['name']}]"
            lines.append(line)

    elif tapping_mode == MODE_SAME_RS:
        rs = rs_entries[0] if rs_entries else None
        if rs and rs.get("load"):
            main = f"Fuse {rs['load']}"
            if rs.get("ampere"):
                main += f" ({rs['ampere']})"
            main += f" : {equipment} Power Tapped – MAIN"
            if rs.get("name"):
                main += f"  [{rs['name']}]"
            lines.append(main)

            if rs.get("rdnt_load"):
                rdnt_amp = rs.get("rdnt_ampere") or rs.get("ampere", "")
                rdnt = f"Fuse {rs['rdnt_load']}"
                if rdnt_amp:
                    rdnt += f" ({rdnt_amp})"
                rdnt += f" : {equipment} Power Tapped – RDNT"
                if rs.get("name"):
                    rdnt += f"  [{rs['name']}]"
                lines.append(rdnt)

    elif tapping_mode == MODE_SEPARATE_RS:
        rs1 = rs_entries[0] if len(rs_entries) > 0 else None
        rs2 = rs_entries[1] if len(rs_entries) > 1 else None

        if rs1 and rs1.get("load"):
            line = f"Fuse {rs1['load']}"
            if rs1.get("ampere"):
                line += f" ({rs1['ampere']})"
            line += f" : {equipment} Power Tapped – MAIN"
            if rs1.get("name"):
                line += f"  [{rs1['name']}]"
            lines.append(line)

        if rs2 and rs2.get("load"):
            line = f"Fuse {rs2['load']}"
            if rs2.get("ampere"):
                line += f" ({rs2['ampere']})"
            line += f" : {equipment} Power Tapped – RDNT"
            if rs2.get("name"):
                line += f"  [{rs2['name']}]"
            lines.append(line)

    return lines


# =============================================================
# RS SECTION PROCESSOR
# =============================================================

def process_rs_section(doc, data: dict, rs_entries: list,
                        tapping_mode: str, total_rectifiers: int,
                        warnings: list):
    olt_label = build_olt_label(data["equipment"], data["olt_label_custom"])
    equipment  = data["equipment"]

    rs1 = rs_entries[0] if len(rs_entries) > 0 else None
    rs2 = rs_entries[1] if len(rs_entries) > 1 else None

    show_rs2_section = (total_rectifiers >= 2)

    # ----------------------------------------------------------
    # Locate template anchors
    # ----------------------------------------------------------
    rect_paras = []
    for p in iter_all_paragraphs(doc):
        if "PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE" in \
                paragraph_full_text(p).upper():
            rect_paras.append(p)

    rs1_fuse_para, _ = find_in_doc(doc, [
        "{{load}}",
        "{{load}}+ Equipment",
        "FUSE No: L3 Nokia OLT MF-2",
        "FUSE No: L3 OLT MF-2",
        "FUSE No: L3",
    ])

    rs2_fuse_para, _ = find_in_doc(doc, [
        "FUSE No: L6 Nokia OLT MF-2",
        "FUSE No: L6 OLT MF-2",
        "FUSE No: L6",
    ])

    # ----------------------------------------------------------
    # RECTIFIER 1 header  (FIX 1: cleaner format)
    # ----------------------------------------------------------
    if rect_paras and rs1:
        if tapping_mode == MODE_SINGLE:
            r1_header = (
                f"PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE: "
                f"(RECTIFIER 1 – {rs1['name']})"
            )
        elif tapping_mode == MODE_SAME_RS:
            rdnt_info = (
                f" / RDNT: {rs1['rdnt_load']}"
                if rs1.get("rdnt_load") else ""
            )
            r1_header = (
                f"PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE: "
                f"(RECTIFIER 1 – {rs1['name']}"
                f" / MAIN: {rs1['load']}{rdnt_info})"
            )
        elif tapping_mode == MODE_SEPARATE_RS:
            r1_header = (
                f"PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE: "
                f"(RECTIFIER 1 – {rs1['name']} – MAIN)"
            )
        set_paragraph_text(rect_paras[0], r1_header)

    # ----------------------------------------------------------
    # RS1 FUSE line(s)
    # FIX 2: Use insert_paragraph_after_copy_format for RDNT line
    # ----------------------------------------------------------
    if rs1_fuse_para and rs1:
        if tapping_mode == MODE_SINGLE:
            set_paragraph_text(
                rs1_fuse_para,
                fuse_no_line(rs1["load"], olt_label, equipment)
            )

        elif tapping_mode == MODE_SAME_RS:
            # MAIN line (replace in-place, keeps original formatting)
            set_paragraph_text(
                rs1_fuse_para,
                fuse_no_line(rs1["load"], olt_label, equipment, "MAIN")
            )
            # RDNT line (copy format from MAIN line)
            if rs1.get("rdnt_load"):
                insert_paragraph_after_copy_format(
                    rs1_fuse_para,
                    fuse_no_line(
                        rs1["rdnt_load"], olt_label, equipment, "RDNT")
                )

        elif tapping_mode == MODE_SEPARATE_RS:
            set_paragraph_text(
                rs1_fuse_para,
                fuse_no_line(rs1["load"], olt_label, equipment, "MAIN")
            )

    # ----------------------------------------------------------
    # RECTIFIER 2 header
    # ----------------------------------------------------------
    if len(rect_paras) >= 2:
        if not show_rs2_section:
            set_paragraph_text(rect_paras[1], "")

        elif tapping_mode == MODE_SEPARATE_RS and rs2:
            set_paragraph_text(
                rect_paras[1],
                f"PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE: "
                f"(RECTIFIER 2 – {rs2['name']} – RDNT)"
            )
        else:
            rs2_display = rs2["name"] if rs2 else "Rectifier 2"
            set_paragraph_text(
                rect_paras[1],
                f"PROPOSED RECTIFIER SYSTEM LOAD BREAKER FUSE: "
                f"(RECTIFIER 2 – {rs2_display} – Not Used for Tapping)"
            )

    # ----------------------------------------------------------
    # RS2 FUSE line
    # ----------------------------------------------------------
    if rs2_fuse_para:
        if tapping_mode == MODE_SEPARATE_RS and rs2 and show_rs2_section:
            set_paragraph_text(
                rs2_fuse_para,
                fuse_no_line(rs2["load"], olt_label, equipment, "RDNT")
            )
        elif show_rs2_section:
            set_paragraph_text(
                rs2_fuse_para,
                "No tapping point assigned to this rectifier."
            )
        else:
            set_paragraph_text(rs2_fuse_para, "")

    # ----------------------------------------------------------
    # RECTIFIER 1 caption
    # ----------------------------------------------------------
    for p in iter_all_paragraphs(doc):
        if paragraph_full_text(p).strip() == "RECTIFIER 1":
            if rs1:
                if tapping_mode == MODE_SAME_RS:
                    rdnt = (
                        f" / RDNT: {rs1['rdnt_load']}"
                        if rs1.get("rdnt_load") else ""
                    )
                    cap = (
                        f"RECTIFIER 1 – {rs1['name']} "
                        f"(MAIN: {rs1['load']}{rdnt})"
                    )
                elif tapping_mode == MODE_SEPARATE_RS:
                    cap = f"RECTIFIER 1 – {rs1['name']} – MAIN"
                else:
                    cap = f"RECTIFIER 1 – {rs1['name']}"
                set_paragraph_text(p, cap)
            break

    # ----------------------------------------------------------
    # RECTIFIER 2 caption
    # ----------------------------------------------------------
    for p in iter_all_paragraphs(doc):
        if paragraph_full_text(p).strip() == "RECTIFIER 2":
            if not show_rs2_section:
                set_paragraph_text(p, "")
            elif tapping_mode == MODE_SEPARATE_RS and rs2:
                set_paragraph_text(
                    p, f"RECTIFIER 2 – {rs2['name']} – RDNT")
            else:
                rs2_display = rs2["name"] if rs2 else "Rectifier 2"
                set_paragraph_text(
                    p,
                    f"RECTIFIER 2 – {rs2_display} – Not Used for Tapping"
                )
            break

    # ----------------------------------------------------------
    # RS1 Load Schedule image
    # ----------------------------------------------------------
    if rs1 and rs1.get("load_img_bytes"):
        m = replace_placeholder_with_image(
            doc, RS1_LOAD_PH, rs1["load_img_bytes"])
        warnings.append(
            f"✅ RS1 Load Schedule inserted (matched: '{m}')." if m
            else f"⚠️ RS1 Load Schedule placeholder not found. "
                 f"Tried: {RS1_LOAD_PH}"
        )
    else:
        clear_placeholders(doc, RS1_LOAD_PH)

    # ----------------------------------------------------------
    # RS2 Load Schedule image
    # ----------------------------------------------------------
    if show_rs2_section and rs2 and rs2.get("load_img_bytes"):
        m = replace_placeholder_with_image(
            doc, RS2_LOAD_PH, rs2["load_img_bytes"])
        warnings.append(
            f"✅ RS2 Load Schedule inserted (matched: '{m}')." if m
            else f"⚠️ RS2 Load Schedule placeholder not found. "
                 f"Tried: {RS2_LOAD_PH}"
        )
    else:
        clear_placeholders(doc, RS2_LOAD_PH)

    # ----------------------------------------------------------
    # RS1 Existing Rectifier image
    # ----------------------------------------------------------
    if rs1 and rs1.get("existing_img_bytes"):
        m = replace_placeholder_with_image(
            doc, RS1_EXIST_PH, rs1["existing_img_bytes"])
        warnings.append(
            f"✅ RS1 Existing image inserted (matched: '{m}')." if m
            else f"⚠️ RS1 EXISTING IMAGE placeholder not found. "
                 f"Tried: {RS1_EXIST_PH}"
        )
    else:
        clear_placeholders(doc, RS1_EXIST_PH)

    # ----------------------------------------------------------
    # RS2 Existing Rectifier image
    # ----------------------------------------------------------
    if show_rs2_section and rs2 and rs2.get("existing_img_bytes"):
        m = replace_placeholder_with_image(
            doc, RS2_EXIST_PH, rs2["existing_img_bytes"])
        warnings.append(
            f"✅ RS2 Existing image inserted (matched: '{m}')." if m
            else f"⚠️ RS2 EXISTING IMAGE placeholder not found. "
                 f"Tried: {RS2_EXIST_PH}"
        )
    else:
        clear_placeholders(doc, RS2_EXIST_PH)

    warnings.append(
        f"✅ RS section processed: mode={tapping_mode}, "
        f"rectifiers on site={total_rectifiers}"
    )


# =============================================================
# PLANNED ACTIVITY FUSE LINE UPDATER
# FIX 3: Copy format from matched paragraph for all inserted lines
# =============================================================

def update_planned_activity_fuse_line(doc, rs_entries: list,
                                       tapping_mode: str,
                                       data: dict, warnings: list):
    equipment  = data["equipment"]
    fuse_lines = build_tapping_summary(tapping_mode, rs_entries, equipment)

    if not fuse_lines:
        warnings.append("⚠️ No fuse lines built for planned activity.")
        return

    matched_paras = []
    for p in iter_all_paragraphs(doc):
        txt = paragraph_full_text(p)
        if any(n in txt for n in [
            "Power Tapped",
            "Nokia Power Tapped",
            "Fuse L9",
            "Fuse L17",
            "Fuse L",
        ]):
            matched_paras.append(p)

    if not matched_paras:
        warnings.append(
            "⚠️ Fuse assignment line not found in planned activity section."
        )
        return

    # Replace first matched paragraph with first fuse line
    # (keeps original paragraph formatting)
    set_paragraph_text(matched_paras[0], fuse_lines[0])
    ref_para = matched_paras[0]

    # Insert extra lines copying format from the first matched paragraph
    last_para = ref_para
    for extra in fuse_lines[1:]:
        last_para = insert_paragraph_after_copy_format(last_para, extra)

    # Clear remaining old matched paragraphs
    for old in matched_paras[1:]:
        set_paragraph_text(old, "")

    warnings.append(
        f"✅ Planned activity fuse lines: {len(fuse_lines)} line(s) written."
    )


# =============================================================
# SUPPORTING DOCUMENTS
# =============================================================

def insert_supporting_documents(doc, tssr_pdf_name: str):
    if not tssr_pdf_name:
        return
    anchor, _ = find_in_doc(
        doc,
        ["SUPPORTING DOCUMENTS", "Supporting Documents",
         "Supporting Document", "Attachments"]
    )
    if not anchor:
        anchor = doc.paragraphs[-1]
    p = anchor._parent.add_paragraph()
    anchor._p.addnext(p._p)
    p.add_run(f"TSSR: {tssr_pdf_name}")


# =============================================================
# GENERATE
# =============================================================

def generate_docx_bytes(data: dict, rs_entries: list,
                         tapping_mode: str, total_rectifiers: int,
                         tssr_pdf_name: str):
    if not os.path.exists(TEMPLATE_FILE):
        raise FileNotFoundError("Template.docx not found in repo root.")

    doc = Document(TEMPLATE_FILE)
    warnings = []

    replace_everywhere(doc, build_replacements(data))

    process_rs_section(
        doc, data, rs_entries, tapping_mode, total_rectifiers, warnings)

    update_planned_activity_fuse_line(
        doc, rs_entries, tapping_mode, data, warnings)

    insert_supporting_documents(doc, tssr_pdf_name)

    audit = audit_remaining_mf2(doc)
    if audit:
        warnings.append(
            "⚠️ Possible remaining MF-2 references:\n"
            + "\n".join(f"  • {line}" for line in audit[:10])
        )
    else:
        warnings.append(
            "✅ No remaining MF-2/Nokia OLT references detected.")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue(), warnings


# =============================================================
# CLIPBOARD HELPER
# =============================================================

def clipboard_image_to_bytes(eval_key: str):
    js = """
    async () => {
      try {
        const items = await navigator.clipboard.read();
        for (const item of items) {
          for (const type of item.types) {
            if (type.startsWith('image/')) {
              const blob = await item.getType(type);
              const dataUrl = await new Promise((resolve) => {
                const reader = new FileReader();
                reader.onload = () => resolve(reader.result);
                reader.readAsDataURL(blob);
              });
              return dataUrl;
            }
          }
        }
        return null;
      } catch (e) { return "ERROR:" + e.toString(); }
    }
    """
    result = streamlit_js_eval(
        js_expressions=js, key=eval_key, want_output=True)
    if not result:
        return None
    if isinstance(result, str) and result.startswith("ERROR:"):
        st.warning(f"Clipboard paste failed: {result}")
        return None
    if (isinstance(result, str)
            and "," in result
            and result.startswith("data:image/")):
        return base64.b64decode(result.split(",", 1)[1])
    return None


def uploaded_to_bytes(uploaded):
    return uploaded.getvalue() if uploaded else None


def image_input_widget(label_upload, label_paste, upload_key,
                        paste_btn_key, session_key, eval_key,
                        preview_caption=""):
    uploaded = st.file_uploader(
        label_upload,
        type=["png", "jpg", "jpeg", "bmp"],
        key=upload_key,
    )
    if st.button(label_paste, key=paste_btn_key):
        st.session_state[session_key] = clipboard_image_to_bytes(eval_key)
    img_bytes = (uploaded_to_bytes(uploaded)
                 or st.session_state.get(session_key))
    if img_bytes:
        st.image(img_bytes, caption=preview_caption, width=340)
    return img_bytes


# =============================================================
# STREAMLIT UI
# =============================================================

st.set_page_config(page_title="MOP Automation", layout="wide")
st.title("MOP Automation")

if not os.path.exists(TEMPLATE_FILE):
    st.error("**Template.docx not found** in repo root.")
    st.stop()

# ── Debug ─────────────────────────────────────────────────────
with st.expander("🔍 Debug: Inspect all text in template", expanded=False):
    if st.button("List all text (body + header + footer + XML)"):
        _doc = Document(TEMPLATE_FILE)
        st.code("\n".join(debug_list_all_text(_doc)), language="text")

st.divider()

# ── General Information ───────────────────────────────────────
st.subheader("General Information")
g1, g2 = st.columns(2)

with g1:
    site_name        = st.text_input(
        "Site Name", placeholder="e.g. CDO-707-HS")
    plaid            = st.text_input(
        "Plaid", placeholder="e.g. MIN1338")
    equipment        = st.text_input(
        "Equipment", value="Nokia Lightspan MF-2")
    olt_label_custom = st.text_input(
        "Custom OLT Label (leave blank if Nokia Lightspan MF-2)",
        placeholder="e.g. OLT MA5800",
    )

with g2:
    prepared_by     = st.text_input(
        "Prepared By", placeholder="e.g. John Carlo Rabanes")
    position        = st.text_input(
        "Position", value="OLT Rollout Engineer")
    target_datetime = st.text_input(
        "Target Date and Time",
        value="May 19- June 19, 2026 10:00AM-6:00PM",
    )

if equipment:
    _parts  = equipment.strip().split()
    _vendor = _parts[0] if _parts else "Nokia"
    _model  = _parts[-1] if len(_parts) > 1 else equipment
    _olt    = (
        "OLT MF-2"
        if normalize_spaces(equipment).lower() == "nokia lightspan mf-2"
        else (olt_label_custom.strip() or equipment)
    )
    st.info(
        f"**Equipment replacement preview:**  \n"
        f"- `Nokia Lightspan MF-2` → `{equipment}`  \n"
        f"- `OLT MF-2` → `{_olt}`  \n"
        f"- `Nokia OLT MF-2` → `{_vendor} {_olt}`  \n"
        f"- `Nokia / OLT MF-2` → `{_vendor} / {_olt}`  \n"
        f"- `MF-2` → `{_model}`"
    )

st.divider()

# ── Power Tapping Configuration ───────────────────────────────
st.subheader("Power Tapping Configuration")
pc1, pc2 = st.columns(2)

with pc1:
    total_rectifiers = st.selectbox(
        "Total number of rectifiers on site",
        options=[1, 2],
        index=1,
        help=(
            "Total physical rectifiers on site. "
            "All rectifiers will be documented regardless of tapping mode."
        ),
    )

with pc2:
    tapping_mode = st.radio(
        "Tapping Point Setup",
        options=[MODE_SINGLE, MODE_SAME_RS, MODE_SEPARATE_RS],
        index=0,
        help=(
            f"**{MODE_SINGLE}**: One fuse – MAIN only.  \n"
            f"**{MODE_SAME_RS}**: MAIN + RDNT from the same rectifier.  \n"
            f"**{MODE_SEPARATE_RS}**: RS1 = MAIN, RS2 = RDNT "
            f"(different rectifiers)."
        ),
    )


def get_scenario_description(mode, total):
    if total == 1:
        scenarios = {
            MODE_SINGLE:      "📌 1 rectifier on site. Single tapping point (MAIN only). No RS2 section.",
            MODE_SAME_RS:     "📌 1 rectifier on site. MAIN + RDNT from same rectifier. No RS2 section.",
            MODE_SEPARATE_RS: "⚠️ Separate Rectifiers mode requires 2 rectifiers. Set total to 2.",
        }
    else:
        scenarios = {
            MODE_SINGLE:      "📌 2 rectifiers on site. RS1 tapped (MAIN). RS2 documented but not tapped.",
            MODE_SAME_RS:     "📌 2 rectifiers on site. RS1 MAIN + RDNT. RS2 documented but not tapped.",
            MODE_SEPARATE_RS: "📌 2 rectifiers on site. RS1 = MAIN tapping. RS2 = RDNT tapping.",
        }
    return scenarios.get(mode, "")


desc = get_scenario_description(tapping_mode, total_rectifiers)
if "⚠️" in desc:
    st.warning(desc)
else:
    st.info(desc)

st.divider()

# ── RS Details ────────────────────────────────────────────────
st.subheader("RS Details")
st.info("ℹ️ **Load Assignment = Fuse No.** e.g. `F8`, `L3`, `L6`")

rs_entries = []

# ── RS1 ───────────────────────────────────────────────────────
rs1_labels = {
    MODE_SINGLE:      "RS1 Details – Tapping Point",
    MODE_SAME_RS:     "RS1 Details – MAIN + RDNT (Same Rectifier)",
    MODE_SEPARATE_RS: "RS1 Details – MAIN Tapping Point",
}
with st.expander(rs1_labels.get(tapping_mode, "RS1 Details"), expanded=True):
    fa, fb = st.columns(2)

    with fa:
        st.markdown("#### RS1 Information")
        rs1_name   = st.text_input(
            "RS1 Rectifier Name",
            placeholder="e.g. Eltek Flatpack 1",
            key="rs1_name",
        )
        rs1_load_label = {
            MODE_SINGLE:      "RS1 Load Assignment (= Fuse No.)",
            MODE_SAME_RS:     "RS1 MAIN Load Assignment (= MAIN Fuse No.)",
            MODE_SEPARATE_RS: "RS1 MAIN Load Assignment (= Fuse No.)",
        }
        rs1_load   = st.text_input(
            rs1_load_label.get(tapping_mode, "RS1 Load Assignment"),
            placeholder="e.g. F8",
            key="rs1_load",
        )
        rs1_ampere = st.text_input(
            "RS1 Breaker Size",
            placeholder="e.g. 10A",
            key="rs1_ampere",
        )

        if tapping_mode == MODE_SAME_RS:
            st.markdown("---")
            st.markdown("**RDNT Tapping Point (same rectifier)**")
            rs1_rdnt_load   = st.text_input(
                "RDNT Load Assignment (= RDNT Fuse No.)",
                placeholder="e.g. F9",
                key="rs1_rdnt_load",
            )
            rs1_rdnt_ampere = st.text_input(
                "RDNT Breaker Size",
                placeholder="e.g. 10A",
                key="rs1_rdnt_ampere",
            )
        else:
            rs1_rdnt_load   = ""
            rs1_rdnt_ampere = ""

    with fb:
        st.markdown("#### RS1 Load Schedule Image")
        st.caption("Fuse panel / load schedule photo")
        rs1_load_img = image_input_widget(
            label_upload    = "Upload RS1 Load Schedule image",
            label_paste     = "Paste RS1 Load Schedule from clipboard",
            upload_key      = "rs1_load_img_upload",
            paste_btn_key   = "paste_rs1_load_btn",
            session_key     = "rs1_load_img_bytes",
            eval_key        = "rs1_load_clip_eval",
            preview_caption = "RS1 Load Schedule",
        )

    st.markdown("#### RS1 Existing Rectifier Photo")
    st.caption("Physical photo of the existing rectifier")
    rs1_exist_img = image_input_widget(
        label_upload    = "Upload RS1 Existing Rectifier image",
        label_paste     = "Paste RS1 Existing Rectifier from clipboard",
        upload_key      = "rs1_existing_img_upload",
        paste_btn_key   = "paste_rs1_existing_btn",
        session_key     = "rs1_existing_img_bytes",
        eval_key        = "rs1_existing_clip_eval",
        preview_caption = "RS1 Existing Rectifier",
    )

rs_entries.append({
    "name":               rs1_name.strip(),
    "load":               rs1_load.strip(),
    "ampere":             rs1_ampere.strip(),
    "rdnt_load":          rs1_rdnt_load.strip(),
    "rdnt_ampere":        rs1_rdnt_ampere.strip(),
    "load_img_bytes":     rs1_load_img,
    "existing_img_bytes": rs1_exist_img,
})

# ── RS2 ───────────────────────────────────────────────────────
if total_rectifiers >= 2:
    rs2_labels = {
        MODE_SINGLE:      "RS2 Details – On Site (Not Used for Tapping)",
        MODE_SAME_RS:     "RS2 Details – On Site (Not Used for Tapping)",
        MODE_SEPARATE_RS: "RS2 Details – RDNT Tapping Point",
    }
    with st.expander(
            rs2_labels.get(tapping_mode, "RS2 Details"), expanded=True):

        if tapping_mode in (MODE_SINGLE, MODE_SAME_RS):
            st.info(
                "ℹ️ This rectifier is on site but **not used for tapping**. "
                "It will be fully documented in the MOP."
            )

        fa, fb = st.columns(2)

        with fa:
            st.markdown("#### RS2 Information")
            rs2_name   = st.text_input(
                "RS2 Rectifier Name",
                placeholder="e.g. Eltek Flatpack 2",
                key="rs2_name",
            )
            if tapping_mode == MODE_SEPARATE_RS:
                rs2_load   = st.text_input(
                    "RS2 RDNT Load Assignment (= Fuse No.)",
                    placeholder="e.g. L6",
                    key="rs2_load",
                )
                rs2_ampere = st.text_input(
                    "RS2 Breaker Size",
                    placeholder="e.g. 10A",
                    key="rs2_ampere",
                )
            else:
                rs2_load   = ""
                rs2_ampere = ""
                st.caption("No tapping assignment for this rectifier.")

        with fb:
            st.markdown("#### RS2 Load Schedule Image")
            st.caption("Fuse panel / load schedule photo (documentation)")
            rs2_load_img = image_input_widget(
                label_upload    = "Upload RS2 Load Schedule image",
                label_paste     = "Paste RS2 Load Schedule from clipboard",
                upload_key      = "rs2_load_img_upload",
                paste_btn_key   = "paste_rs2_load_btn",
                session_key     = "rs2_load_img_bytes",
                eval_key        = "rs2_load_clip_eval",
                preview_caption = "RS2 Load Schedule",
            )

        st.markdown("#### RS2 Existing Rectifier Photo")
        st.caption("Physical photo of the existing rectifier (documentation)")
        rs2_exist_img = image_input_widget(
            label_upload    = "Upload RS2 Existing Rectifier image",
            label_paste     = "Paste RS2 Existing Rectifier from clipboard",
            upload_key      = "rs2_existing_img_upload",
            paste_btn_key   = "paste_rs2_existing_btn",
            session_key     = "rs2_existing_img_bytes",
            eval_key        = "rs2_existing_clip_eval",
            preview_caption = "RS2 Existing Rectifier",
        )

    rs_entries.append({
        "name":               rs2_name.strip(),
        "load":               rs2_load.strip(),
        "ampere":             rs2_ampere.strip(),
        "rdnt_load":          "",
        "rdnt_ampere":        "",
        "load_img_bytes":     rs2_load_img,
        "existing_img_bytes": rs2_exist_img,
    })

# ── Planned activity preview ──────────────────────────────────
_eq = equipment.strip() or "Equipment"
preview_lines = build_tapping_summary(tapping_mode, rs_entries, _eq)
if any(rs.get("load") for rs in rs_entries) and preview_lines:
    st.divider()
    st.markdown("**📋 Planned Activity Fuse Lines Preview:**")
    for pl in preview_lines:
        st.code(pl)

st.divider()

# ── Supporting Documents ──────────────────────────────────────
st.subheader("Supporting Documents")
tssr_pdf = st.file_uploader(
    "TSSR PDF (filename will be written into the Word document)",
    type=["pdf"], key="tssr_pdf",
)
if tssr_pdf:
    st.success(f"PDF ready: {tssr_pdf.name}")

st.divider()

# ── Generate ──────────────────────────────────────────────────
data = {
    "site_name":        site_name.strip(),
    "plaid":            plaid.strip(),
    "equipment":        equipment.strip(),
    "olt_label_custom": olt_label_custom.strip(),
    "prepared_by":      prepared_by.strip(),
    "position":         position.strip(),
    "target_datetime":  target_datetime.strip(),
}

required_base = [
    "site_name", "plaid", "equipment",
    "prepared_by", "position", "target_datetime",
]

if st.button("Generate MOP (.docx)", type="primary"):

    missing = [k for k in required_base if not data.get(k)]
    if missing:
        st.error("Please fill in: " + ", ".join(missing))
        st.stop()

    rs_valid = True

    if not rs_entries[0].get("name"):
        st.error("RS1 Rectifier Name is required.")
        rs_valid = False
    if not rs_entries[0].get("load"):
        st.error("RS1 Load Assignment is required.")
        rs_valid = False
    if tapping_mode == MODE_SAME_RS and not rs_entries[0].get("rdnt_load"):
        st.error("RDNT Load Assignment is required for Same Rectifier mode.")
        rs_valid = False
    if total_rectifiers >= 2:
        if len(rs_entries) < 2 or not rs_entries[1].get("name"):
            st.error("RS2 Rectifier Name is required when 2 rectifiers are on site.")
            rs_valid = False
    if tapping_mode == MODE_SEPARATE_RS:
        if len(rs_entries) < 2 or not rs_entries[1].get("load"):
            st.error("RS2 Load Assignment is required for Separate Rectifiers mode.")
            rs_valid = False
        if total_rectifiers < 2:
            st.error("Separate Rectifiers mode requires 2 rectifiers on site.")
            rs_valid = False

    if not rs_valid:
        st.stop()

    try:
        docx_bytes, gen_warnings = generate_docx_bytes(
            data=data,
            rs_entries=rs_entries,
            tapping_mode=tapping_mode,
            total_rectifiers=total_rectifiers,
            tssr_pdf_name=(tssr_pdf.name if tssr_pdf else ""),
        )

        out_name = (
            f"MOP_{safe_filename(data['site_name'])}"
            f"_{safe_filename(data['plaid'])}.docx"
        )

        st.success("MOP generated successfully!")
        for w in gen_warnings:
            if w.startswith("✅"):
                st.success(w)
            elif w.startswith("⚠️"):
                st.warning(w)
            else:
                st.info(w)

        st.download_button(
            label="⬇️ Download MOP",
            data=docx_bytes,
            file_name=out_name,
            mime=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )

    except Exception as e:
        st.exception(e)